from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from typing import Any, Dict, List
from datetime import datetime
from ...base import BaseExportService


class DOCXExportService(BaseExportService):
    """Servicio para generar reportes DOCX con formato institucional y sobrio"""

    FIELDS = [
        ("id", "ID de reporte"),
        ("userId", "ID de usuario"),
        ("user.documentId", "Documento de usuario"),
        ("user.userInformation.name", "Nombre del usuario"),
        ("user.userInformation.lastName", "Apellido del usuario"),
        ("externalNameUser", "Nombre externo del usuario"),
        ("externalOrganization", "Organizacion externa"),
        ("reportTitle", "Titulo del reporte"),
        ("conversation", "Conversacion"),
        ("base", "Base"),
        ("createdAt", "Fecha de creacion"),
        ("updatedAt", "Fecha de actualizacion"),
        ("unity", "Unidad"),
        ("rig", "Equipo (rig)"),
        ("project", "Proyecto"),
        ("field", "Campo"),
        ("reportType", "Tipo de reporte"),
        ("hazardClassification", "Clasificacion del peligro"),
        ("hazardType", "Tipo de peligro"),
        ("detailedDescription", "Descripcion detallada"),
        ("findingCause", "Causa del hallazgo"),
        ("reportEvidence", "Evidencias del reporte"),
        ("actions", "Acciones"),
        ("reportStatus", "Estado del reporte"),
        ("loraReportCode", "Codigo de reporte LORA"),
    ]
    def __init__(self):
        self.document = None

    def generate_file(self, data: Any, options: Dict = None) -> io.BytesIO:
        if not self.validate_data(data):
            raise ValueError("Datos no válidos")

        self.document = Document()
        self._apply_base_styles()

        # Secciones principales del documento
        self._render_header(data)
        self._render_fields_table(data)
        self._render_actions(data.get("actions", []))
        self._render_footer(data)

        # Guardar documento en buffer
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)

        # Log de control
        print("DOCX buffer size:", buffer.getbuffer().nbytes)

        return buffer

    # ---------------------------
    # RENDERIZADORES PRINCIPALES
    # ---------------------------

    def _render_header(self, data: Dict):
        """Encabezado institucional"""
        title = data.get("reportTitle", "Reporte de Hallazgo").upper()
        code = data.get("loraReportCode", "N/A")
        status = data.get("reportStatus", "N/A").upper()
        created = data.get("createdAt", "N/A")

        # Título principal
        heading = self.document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtítulo con código y estado
        p = self.document.add_paragraph(f"Código: {code}    Estado: {status}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Fecha de creación
        p2 = self.document.add_paragraph(f"Fecha de creación: {created}")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self._add_separator()

    def _render_fields_table(self, data: Dict):
        """Tabla con todos los campos alineados a XLSX all_reports"""
        self.document.add_heading("Detalles del reporte", level=1)

        table = self.document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Campo"
        hdr_cells[1].text = "Valor"

        for key, label in self.FIELDS:
            if key == "actions":
                value = self._format_actions(data.get("actions", []))
            else:
                value = self._format_value(self._get_nested_value(data, key, fallback=data))
            row = table.add_row().cells
            row[0].text = str(label)
            row[1].text = str(value if value not in (None, "") else "N/A")

        self._add_separator()

    def _render_actions(self, actions: List[Dict]):
        """Sección de acciones"""
        self.document.add_heading("Acciones", level=1)
        if not actions:
            self.document.add_paragraph("Sin acciones registradas.")
            return

        for i, act in enumerate(actions, 1):
            desc = act.get("description", "N/A")
            resp = act.get("responsible", "No asignado")
            due = act.get("dueDate", "N/A")
            status = act.get("status", "N/A").upper()

            self.document.add_paragraph(f"{i}. {desc}", style="List Number")
            self.document.add_paragraph(f"   Responsable: {resp}")
            self.document.add_paragraph(f"   Fecha límite: {due}")
            self.document.add_paragraph(f"   Estado: {status}")
            self.document.add_paragraph()

        self._add_separator()

    def _render_footer(self, data: Dict):
        """Pie de documento con identificación"""
        section = self.document.sections[-1]
        footer = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = (
            f"ID: {data.get('id', 'N/A')}  |  "
            f"Exportado: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  "
            "VALERA ECOSYSTEM"
        )

    # ---------------------------
    # UTILITARIOS / ESTILO
    # ---------------------------

    def _get_nested_value(self, obj: Dict, key: str, fallback: Dict = None) -> Any:
        """Navega claves anidadas usando punto. Permite fallback para casos especiales."""
        parts = key.split('.')
        value = obj
        for part in parts:
            if isinstance(value, dict):
                value = value.get(part, None)
            else:
                value = None
            if value is None:
                break
        if value is None and fallback and key == "reportEvidence":
            evidences = fallback.get("evidence", [])
            return evidences if evidences else None
        if value is None and fallback and key == "actions":
            return fallback.get("actions", [])
        return value

    def _format_value(self, value: Any) -> Any:
        """Normaliza listas/diccionarios a texto legible."""
        if value is None:
            return "N/A"
        if isinstance(value, list):
            formatted_items = []
            for item in value:
                if isinstance(item, dict):
                    parts = [f"{k}: {v}" for k, v in item.items() if v not in (None, "")]
                    formatted_items.append(", ".join(parts) if parts else str(item))
                else:
                    formatted_items.append(str(item))
            return "\n".join(formatted_items) if formatted_items else "N/A"
        if isinstance(value, dict):
            parts = [f"{k}: {v}" for k, v in value.items() if v not in (None, "")]
            return ", ".join(parts) if parts else "N/A"
        return value

    def _apply_base_styles(self):
        """Estilos base para texto institucional"""
        style = self.document.styles["Normal"]
        font = style.font
        font.name = "Courier New"
        font.size = Pt(10)

    def _add_separator(self):
        """Agrega un pequeño espacio como separador visual"""
        self.document.add_paragraph("\n")

    def get_content_type(self) -> str:
        """Retorna el tipo MIME del DOCX"""
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def get_file_extension(self) -> str:
        """Retorna la extensión de archivo DOCX"""
        return ".docx"
