from docx import Document
from docx.shared import Inches
import io
from typing import Any, Dict, List
from ..services.base import BaseExportService


class DOCXExportService(BaseExportService):
    """Servicio para generar archivos DOCX"""
    
    def __init__(self):
        self.document = None
    
    def generate_file(self, data: Any, options: Dict = None) -> io.BytesIO:
        """Genera un archivo DOCX"""
        if not self.validate_data(data):
            raise ValueError("Los datos proporcionados no son válidos")
        
        self.document = Document()
        
        # Si los datos son un diccionario con estructura de documento
        if isinstance(data, dict):
            self._process_document_data(data, options or {})
        else:
            # Datos simples como texto
            self._add_simple_text(str(data))
        
        # Guardar en buffer
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer
    
    def _process_document_data(self, data: Dict, options: Dict):
        """Procesa datos estructurados de documento"""
        # Título del documento
        if 'title' in data and data['title']:
            title = self.document.add_heading(data['title'], 0)
            title.alignment = 1  # Centrado
        
        # Contenido de párrafos
        if 'content' in data and isinstance(data['content'], list):
            for paragraph_text in data['content']:
                self.document.add_paragraph(paragraph_text)
        
        # Tablas
        if 'tables' in data and isinstance(data['tables'], list):
            for table in data['tables']:
                self._add_table(table)
        
        # Si es una tabla simple
        if 'headers' in data and 'rows' in data:
            self._add_table(data)
    
    def _add_table(self, table_data: Dict):
        """Añade una tabla al documento DOCX"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        title = table_data.get('title', '')
        
        if title:
            self.document.add_heading(title, level=2)
        
        if not headers or not rows:
            return
        
        # Crear tabla
        table = self.document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Añadir encabezados
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
            # Hacer los encabezados en negrita
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        # Añadir filas de datos
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                if i < len(row_cells):
                    row_cells[i].text = str(cell_data)
        
        # Añadir espacio después de la tabla
        self.document.add_paragraph()
    
    def _add_simple_text(self, text: str):
        """Añade texto simple al documento"""
        self.document.add_paragraph(text)
    
    def get_content_type(self) -> str:
        """Retorna el tipo de contenido MIME para DOCX"""
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    def get_file_extension(self) -> str:
        """Retorna la extensión del archivo DOCX"""
        return ".docx"